"""Office 文档解析器集合：DOCX, DOC, XLSX, XLS, PPT。"""

from __future__ import annotations

import base64
import io
import logging
import os
import re
import subprocess
import tempfile
import threading
import time
import traceback
import zipfile
from collections.abc import MutableSequence
from concurrent.futures import Future, ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field
from io import BytesIO
from multiprocessing import Manager
from pathlib import Path
from types import TracebackType
from typing import Any, Callable, Dict, Iterable, List, Literal, Optional, Set, Tuple, cast

import pandas as pd  # type: ignore[import-untyped]  # pandas 未安装 stub（pandas-stubs），模块整体按 Any 处理
from docx import Document
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.opc.oxml import parse_xml
from docx.opc.pkgreader import _SerializedRelationship, _SerializedRelationships
from PIL import Image

from aion_knowledge.common.config import settings
from aion_knowledge.common.uuid7 import uuid7
from aion_knowledge.pipeline.parser.base import BaseParser, ParsedDocument
from aion_knowledge.pipeline.parser.chain import FirstParser
from aion_knowledge.pipeline.parser.utils import endecode

logger = logging.getLogger(__name__)


# ============================================================
# python-docx 猴子补丁：修复 NULL 关系加载
#（补丁来自 https://github.com/python-openxml/python-docx/issues/1105）
# ============================================================
def load_from_xml_v2(baseURI: str, rels_item_xml: str | None) -> _SerializedRelationships:  # noqa: N803
    """
    返回加载了 *rels_item_xml* 中所含关系的 |_SerializedRelationships| 实例。
    如果 *rels_item_xml* 为 |None|，则返回空集合。
    """
    srels = _SerializedRelationships()  # type: ignore[no-untyped-call]  # docx 第三方 __init__ 未注解，无法在本模块修复
    if rels_item_xml is not None:
        rels_elm: Any = parse_xml(rels_item_xml)  # docx 标注为通用 lxml 元素，实为 CT_Relationships
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ("../NULL", "NULL"):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))  # type: ignore[no-untyped-call]  # docx 第三方构造函数未注解
    return srels


_SerializedRelationships.load_from_xml = load_from_xml_v2  # type: ignore[method-assign]  # 猴子补丁替换第三方方法（python-docx issue #1105 修复）


# ============================================================
# 临时文件/目录上下文管理器
# ============================================================
class TempFileContext:
    """临时文件上下文管理器，退出时自动删除。"""

    def __init__(self, file_content: bytes, suffix: str):
        self.file_content = file_content
        self.suffix = suffix
        self.temp_file: tempfile._TemporaryFileWrapper[bytes] | None = None

    def __enter__(self) -> str:
        self.temp_file = tempfile.NamedTemporaryFile(
            suffix=self.suffix, delete=False
        )
        self.temp_file.write(self.file_content)
        self.temp_file.flush()
        logger.info(
            "Saved %s content to temporary file: %s",
            self.suffix,
            self.temp_file.name,
        )
        return self.temp_file.name

    def __exit__(
        self,
        exc_type: type[BaseException] | None,
        exc_val: BaseException | None,
        exc_tb: TracebackType | None,
    ) -> Literal[False]:  # 永远不吞异常：返回 False（Literal 供 mypy 正确分析 with 块控制流）
        if self.temp_file:
            self.temp_file.close()
            if os.path.exists(self.temp_file.name):
                os.remove(self.temp_file.name)
            logger.info("File %s has been deleted.", self.temp_file.name)
        return False


class TempDirContext:
    """临时目录上下文管理器，退出时自动删除。"""

    def __init__(self) -> None:
        self.temp_dir: tempfile.TemporaryDirectory[str] | None = None

    def __enter__(self) -> str:
        self.temp_dir = tempfile.TemporaryDirectory()
        logger.info("Created temporary directory: %s", self.temp_dir.name)
        return self.temp_dir.name

    def __exit__(
        self,
        exc_type: type[BaseException] | None,
        exc_val: BaseException | None,
        exc_tb: TracebackType | None,
    ) -> Literal[False]:  # 永远不吞异常：返回 False（Literal 供 mypy 正确分析 with 块控制流）
        if self.temp_dir and os.path.exists(self.temp_dir.name):
            self.temp_dir.cleanup()
            logger.info("Directory %s has been deleted.", self.temp_dir.name)
        return False


# ============================================================
# XLSX 修复工具（来源：xlsx_repair.py）
# ============================================================
SST_PART = "xl/sharedStrings.xml"
_SST_OVERRIDE_RE = re.compile(
    r'<Override[^>]*PartName="[^"]*sharedStrings\.xml"[^>]*/>',
    re.IGNORECASE,
)
_SST_REL_RE = re.compile(
    r'<Relationship[^>]*Type="[^"]*sharedStrings"[^>]*/>',
    re.IGNORECASE,
)


def repair_xlsx_bytes(content: bytes) -> bytes | None:
    """返回修复后的 XLSX 字节，如果未应用修复则返回 None。

    处理在包元数据中引用 ``xl/sharedStrings.xml`` 但缺少该部件的工
    作簿（某些导出工具的常见情况）。当工作表仅使用内联字符串时，
    清单引用会被剥离，以便 openpyxl 能够读取文件。
    """
    if not zipfile.is_zipfile(io.BytesIO(content)):
        return None

    with zipfile.ZipFile(io.BytesIO(content), "r") as zin:
        names = _normalized_names(zin.namelist())
        sst_path = _find_shared_strings_path(names)
        if sst_path:
            if sst_path == SST_PART:
                return None
            return _rewrite_zip(
                zin, lambda files: _rename_shared_strings_part(files, sst_path)
            )
        if not _package_references_shared_strings(zin, names):
            return None
        if _worksheets_use_shared_string_cells(zin, names):
            return None
        return _rewrite_zip(zin, _strip_shared_strings_manifest)


def _normalized_names(namelist: Iterable[str]) -> Set[str]:
    return {name.replace("\\", "/") for name in namelist}


def _find_shared_strings_path(names: Set[str]) -> str | None:
    for name in names:
        if name.lower().endswith("sharedstrings.xml"):
            return name
    return None


def _package_references_shared_strings(
    zin: zipfile.ZipFile, names: Set[str]
) -> bool:
    content_types = "[Content_Types].xml"
    if content_types in names:
        ct = zin.read(content_types).decode("utf-8", errors="replace")
        if "sharedstrings.xml" in ct.lower():
            return True

    rels_path = "xl/_rels/workbook.xml.rels"
    if rels_path in names:
        rels = zin.read(rels_path).decode("utf-8", errors="replace")
        if "sharedstrings" in rels.lower():
            return True
    return False


def _worksheets_use_shared_string_cells(
    zin: zipfile.ZipFile, names: Set[str]
) -> bool:
    for name in names:
        if not name.startswith("xl/worksheets/") or not name.endswith(".xml"):
            continue
        sheet = zin.read(name).decode("utf-8", errors="replace")
        if re.search(r'\bt="s"', sheet):
            return True
    return False


def _rename_shared_strings_part(
    files: Dict[str, bytes], source_path: str
) -> Dict[str, bytes]:
    updated = dict(files)
    updated[SST_PART] = updated.pop(source_path)
    return updated


def _strip_shared_strings_manifest(files: Dict[str, bytes]) -> Dict[str, bytes]:
    updated = dict(files)
    ct_path = "[Content_Types].xml"
    if ct_path in updated:
        ct = updated[ct_path].decode("utf-8")
        ct = _SST_OVERRIDE_RE.sub("", ct)
        updated[ct_path] = ct.encode("utf-8")

    rels_path = "xl/_rels/workbook.xml.rels"
    if rels_path in updated:
        rels = updated[rels_path].decode("utf-8")
        rels = _SST_REL_RE.sub("", rels)
        updated[rels_path] = rels.encode("utf-8")
    return updated


def _rewrite_zip(
    zin: zipfile.ZipFile,
    transform: Callable[[Dict[str, bytes]], Dict[str, bytes]],
) -> bytes:
    files: Dict[str, bytes] = {}
    for info in zin.infolist():
        name = info.filename.replace("\\", "/")
        files[name] = zin.read(info.filename)
    files = transform(files)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    return out.getvalue()


# ============================================================
# XLSX 合并单元格工具（来源：xlsx_merge.py）
# ============================================================
def fill_merged_cells_xlsx(content: bytes) -> bytes:
    """取消合并范围并将主单元格值复制到所有被覆盖的单元格。

    openpyxl 仅在合并的左上角单元格中存储值；pandas 在其他单元格中看到 NaN。
    填充后可以使逐行的 RAG 块保留上下文。
    """
    if not zipfile.is_zipfile(BytesIO(content)):
        return content

    from openpyxl import (  # type: ignore[import-untyped]  # openpyxl 未安装 stub（types-openpyxl）
        load_workbook,
    )

    wb = load_workbook(BytesIO(content), data_only=True)
    changed = False
    for ws in wb.worksheets:
        if not ws.merged_cells.ranges:
            continue
        for merge_range in list(ws.merged_cells.ranges):
            master_value = ws.cell(merge_range.min_row, merge_range.min_col).value
            ws.unmerge_cells(str(merge_range))
            for row in range(merge_range.min_row, merge_range.max_row + 1):
                for col in range(merge_range.min_col, merge_range.max_col + 1):
                    ws.cell(row, col).value = master_value
            changed = True

    if not changed:
        return content

    out = BytesIO()
    wb.save(out)
    logger.info("Filled merged cells in XLSX before parse")
    return out.getvalue()


# ============================================================
# Excel 转换工具（来源：excel_convert.py）
# ============================================================
_XLS_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_ZIP_MAGIC = b"PK\x03\x04"


def detect_excel_format(content: bytes) -> str | None:
    """返回 pandas/excel 格式标识：xlsx、xls、xlsb、ods 或 None。"""
    if not content:
        return None

    from pandas.io.excel._base import (  # type: ignore[import-untyped]  # pandas 未安装 stub
        inspect_excel_format,
    )

    ext = inspect_excel_format(content_or_path=io.BytesIO(content))
    if ext in ("xlsx", "xls", "xlsb", "ods"):
        return cast(str, ext)  # pandas 无 stub 返回 Any，显式断言为 str
    if ext == "zip":
        return "xlsx"

    if content.startswith(_ZIP_MAGIC):
        return "xlsx"
    if len(content) >= len(_XLS_MAGIC) and content.startswith(_XLS_MAGIC):
        return "xls"
    return None


def engine_for_format(ext: str | None) -> str:
    if ext == "xls":
        return "xlrd"
    if ext in ("xlsx", "xlsb"):
        return "openpyxl"
    if ext == "ods":
        return "odf"
    return "openpyxl"


def convert_excel_to_xlsx_bytes(
    content: bytes, suffix: str = ".xlsx"
) -> bytes | None:
    """使用 LibreOffice 将任意电子表格字节转换为 XLSX（如果可用）。"""
    soffice = find_soffice()
    if not soffice:
        return None

    max_attempts = 3
    for attempt in range(1, max_attempts + 1):
        with tempfile.TemporaryDirectory() as temp_dir, \
             tempfile.TemporaryDirectory() as profile_dir:
            src = os.path.join(temp_dir, f"input{suffix}")
            with open(src, "wb") as handle:
                handle.write(content)

            user_installation = Path(profile_dir).as_uri()
            cmd = [
                soffice,
                "--headless",
                f"-env:UserInstallation={user_installation}",
                "--convert-to",
                "xlsx",
                "--outdir",
                temp_dir,
                src,
            ]
            try:
                result = subprocess.run(cmd, capture_output=True, timeout=120)
            except (OSError, subprocess.TimeoutExpired) as exc:
                logger.warning("LibreOffice convert failed to start: %s", exc)
                return None

            if result.returncode != 0:
                stderr = result.stderr.decode("utf-8", errors="ignore")
                logger.warning(
                    "LibreOffice convert failed (attempt %s/%s): %s",
                    attempt,
                    max_attempts,
                    stderr,
                )
                if attempt < max_attempts:
                    time.sleep(0.5 * attempt)
                    continue
                return None

            for name in os.listdir(temp_dir):
                if name.endswith(".xlsx"):
                    with open(os.path.join(temp_dir, name), "rb") as handle:
                        converted = handle.read()
                    logger.info(
                        "Converted spreadsheet via LibreOffice (%s -> xlsx, %d bytes)",
                        suffix,
                        len(converted),
                    )
                    return converted

            if attempt < max_attempts:
                time.sleep(0.5 * attempt)
    return None


def normalize_excel_bytes(
    content: bytes, file_type: str | None = None
) -> bytes:
    """返回 pandas 可读的字节，必要时通过 LibreOffice 转换。"""
    ext = detect_excel_format(content)
    if ext is not None:
        return content

    suffixes = []
    if file_type:
        suffixes.append(f".{file_type.lstrip('.')}")
    suffixes.extend([".xlsx", ".xls", ".et", ".csv"])
    seen: set[str] = set()
    for suffix in suffixes:
        if suffix in seen:
            continue
        seen.add(suffix)
        converted = convert_excel_to_xlsx_bytes(content, suffix=suffix)
        if converted and detect_excel_format(converted) is not None:
            return converted

    raise ValueError(
        "Unrecognized Excel file format; the file may be corrupt, encrypted, "
        "or not a spreadsheet"
    )


def find_soffice() -> Optional[str]:
    possible_paths = [
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        "/opt/libreoffice25.2/program/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]
    for path in possible_paths:
        if path and os.path.exists(path):
            return path

    result = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout.strip()
    return None


# ============================================================
# PPT 转换工具（来源：ppt_convert.py）
# ============================================================
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def is_ole_compound(content: bytes) -> bool:
    return len(content) >= len(_OLE_MAGIC) and content.startswith(_OLE_MAGIC)


def is_zip_openxml(content: bytes) -> bool:
    return len(content) >= len(_ZIP_MAGIC) and content.startswith(_ZIP_MAGIC)


def needs_ppt_to_pptx_conversion(content: bytes, file_type: str | None) -> bool:
    """当内容为旧版 .ppt（OLE）而非现代 .pptx（ZIP）时返回 True。"""
    ext = (file_type or "").lstrip(".").lower()
    if ext == "pptx" or is_zip_openxml(content):
        return False
    if ext == "ppt" or is_ole_compound(content):
        return is_ole_compound(content) or ext == "ppt"
    return False


def convert_ppt_to_pptx_bytes(
    content: bytes, suffix: str = ".ppt"
) -> bytes | None:
    """使用 LibreOffice 将旧版 PowerPoint 字节转换为 PPTX（如果可用）。"""
    soffice = find_soffice()
    if not soffice:
        return None

    max_attempts = 3
    for attempt in range(1, max_attempts + 1):
        with tempfile.TemporaryDirectory() as temp_dir, \
             tempfile.TemporaryDirectory() as profile_dir:
            src = os.path.join(temp_dir, f"input{suffix}")
            with open(src, "wb") as handle:
                handle.write(content)

            user_installation = Path(profile_dir).as_uri()
            cmd = [
                soffice,
                "--headless",
                f"-env:UserInstallation={user_installation}",
                "--convert-to",
                "pptx",
                "--outdir",
                temp_dir,
                src,
            ]
            try:
                result = subprocess.run(cmd, capture_output=True, timeout=120)
            except (OSError, subprocess.TimeoutExpired) as exc:
                logger.warning(
                    "LibreOffice PPT convert failed to start: %s", exc
                )
                return None

            if result.returncode != 0:
                stderr = result.stderr.decode("utf-8", errors="ignore")
                logger.warning(
                    "LibreOffice PPT convert failed (attempt %s/%s): %s",
                    attempt,
                    max_attempts,
                    stderr,
                )
                if attempt < max_attempts:
                    time.sleep(0.5 * attempt)
                    continue
                return None

            for name in os.listdir(temp_dir):
                if name.endswith(".pptx"):
                    with open(os.path.join(temp_dir, name), "rb") as handle:
                        converted = handle.read()
                    logger.info(
                        "Converted presentation via LibreOffice (%s -> pptx, %d bytes)",
                        suffix,
                        len(converted),
                    )
                    return converted

            if attempt < max_attempts:
                time.sleep(0.5 * attempt)
    return None


def normalize_ppt_bytes(
    content: bytes, file_type: str | None
) -> tuple[bytes, str]:
    """返回适合 MarkItDown 的（字节，扩展名）（转换后为 pptx）。"""
    ext = (file_type or "").lstrip(".").lower()

    if is_zip_openxml(content):
        return content, ".pptx"

    if not needs_ppt_to_pptx_conversion(content, ext):
        dotted = f".{ext}" if ext else ".pptx"
        return content, dotted

    suffix = ".ppt" if ext in ("", "ppt") else f".{ext}"
    converted = convert_ppt_to_pptx_bytes(content, suffix=suffix)
    if converted:
        return converted, ".pptx"

    raise ValueError(
        "Legacy PowerPoint (.ppt) is not supported by MarkItDown directly; "
        "LibreOffice is required to convert it to .pptx.  Install LibreOffice "
        "(soffice) in the docreader environment or upload .pptx instead."
    )


# ============================================================
# Docx2Parser — 轻量 DOCX 解析器 (src: docx2_parser.py)
# ============================================================
class Docx2Parser(FirstParser):
    """轻量 DOCX 解析器：优先使用 MarkItDown，失败则回退到完整 DOCX 解析。"""

    # _parser_cls 在 __init__ 中延迟设置（延迟导入避免与 markitdown.py 循环依赖，
    # 且 DocxParser 在本文件中定义较后，需运行时解析）
    def __init__(self, *args: Any, **kwargs: Any) -> None:
        from aion_knowledge.pipeline.parser.markitdown import MarkitdownParser

        self._parser_cls = (MarkitdownParser, DocxParser)
        super().__init__(*args, **kwargs)


# ============================================================
# DocParser — 旧版 .doc 格式解析器（来源：doc_parser.py）
# ============================================================
class SandboxExecutor:
    """带代理配置的沙箱执行器，用于运行命令"""

    def __init__(
        self, proxy: Optional[str] = None, default_timeout: int = 60
    ):
        """初始化沙箱执行器

        参数：
            proxy: 用于网络访问的代理 URL。如果为 None，将使用
                WEB_PROXY 环境变量
            default_timeout: 命令执行的默认超时时间（秒）
        """
        self.proxy = (
            proxy
            or settings.external_https_proxy
            or "http://128.0.0.1:1"
        )
        self.default_timeout = default_timeout

    def execute_in_sandbox(self, cmd: List[str]) -> tuple[bytes, bytes, int]:
        """在沙箱中使用代理配置执行命令

        参数：
            cmd: 要执行的命令

        返回：
            (stdout, stderr, returncode) 元组
        """
        sandbox_methods = [
            self._execute_with_proxy,
        ]

        for method in sandbox_methods:
            try:
                return method(cmd)
            except Exception as e:
                logger.warning(
                    "Sandbox method %s failed: %s", method.__name__, e
                )
                continue

        raise RuntimeError("All sandbox methods failed")

    def _execute_with_proxy(self, cmd: List[str]) -> tuple[bytes, bytes, int]:
        """使用代理配置执行命令"""
        env = os.environ.copy()
        if self.proxy:
            env["http_proxy"] = self.proxy
            env["https_proxy"] = self.proxy
            env["HTTP_PROXY"] = self.proxy
            env["HTTPS_PROXY"] = self.proxy

        logger.info("Executing command with proxy: %s", " ".join(cmd))
        if self.proxy:
            logger.info("Using proxy: %s", self.proxy)

        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            env=env,
        )

        try:
            stdout, stderr = process.communicate(
                timeout=self.default_timeout
            )
            # typeshed 标注 communicate 返回 Optional，PIPE 模式下运行时必为 bytes
            return cast(tuple[bytes, bytes, int], (stdout, stderr, process.returncode))
        except subprocess.TimeoutExpired:
            process.kill()
            raise RuntimeError(
                "Command execution timeout after %s seconds"
                % self.default_timeout
            )


class DocParser(Docx2Parser):
    """DOC 文档解析器"""

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        """使用沙箱执行器初始化 DOC 解析器"""
        super().__init__(*args, **kwargs)
        self.sandbox_executor = SandboxExecutor()

    def parse_into_text(self, content: bytes) -> ParsedDocument:
        logger.info(
            "Parsing DOC document, content size: %s bytes", len(content)
        )

        handle_chain = [
            self._parse_with_docx,
            self._parse_with_antiword,
        ]

        with TempFileContext(content, ".doc") as temp_file_path:
            for handle in handle_chain:
                try:
                    document = handle(temp_file_path)
                    if document:
                        return document
                except Exception as e:
                    logger.warning(
                        "Failed to parse DOC with %s %s",
                        handle.__name__,
                        e,
                    )

            return ParsedDocument(content="")

    def _parse_with_docx(self, temp_file_path: str) -> ParsedDocument:
        logger.info("Multimodal enabled, attempting to extract images from DOC")

        docx_content = self._try_convert_doc_to_docx(temp_file_path)
        if not docx_content:
            raise RuntimeError("Failed to convert DOC to DOCX")

        logger.info("Successfully converted DOC to DOCX, using DocxParser")
        document = super(Docx2Parser, self).parse_into_text(docx_content)
        logger.info(
            "Extracted %s characters using DocxParser", len(document.content)
        )
        return document

    def _parse_with_antiword(self, temp_file_path: str) -> ParsedDocument:
        logger.info("Attempting to parse DOC file with antiword")

        antiword_path = self._try_find_antiword()
        if not antiword_path:
            raise RuntimeError("antiword not found in PATH")

        cmd = [antiword_path, temp_file_path]
        logger.info(
            "Executing antiword in sandbox with proxy configuration"
        )

        stdout, stderr, returncode = self.sandbox_executor.execute_in_sandbox(
            cmd
        )

        if returncode != 0:
            raise RuntimeError(
                "antiword extraction failed: %s"
                % stderr.decode("utf-8", errors="ignore")
            )
        text = stdout.decode("utf-8", errors="ignore")
        logger.info(
            "Successfully extracted %s characters using antiword", len(text)
        )
        return ParsedDocument(content=text)

    def _try_convert_doc_to_docx(
        self, doc_path: str
    ) -> Optional[bytes]:
        """使用 LibreOffice/OpenOffice 将 DOC 文件转换为 DOCX 格式。"""
        logger.info("Converting DOC to DOCX: %s", doc_path)

        soffice_path = self._try_find_soffice()
        if not soffice_path:
            return None

        logger.info("Using %s to convert DOC to DOCX", soffice_path)

        max_attempts = 3
        for attempt in range(1, max_attempts + 1):
            with TempDirContext() as temp_dir, TempDirContext() as profile_dir:
                user_installation = Path(profile_dir).as_uri()
                cmd = [
                    soffice_path,
                    "--headless",
                    "-env:UserInstallation=%s" % user_installation,
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    doc_path,
                ]
                logger.info(
                    "Running command in sandbox (attempt %s/%s): %s",
                    attempt,
                    max_attempts,
                    " ".join(cmd),
                )

                stdout, stderr, returncode = (
                    self.sandbox_executor.execute_in_sandbox(cmd)
                )

                if returncode != 0:
                    logger.warning(
                        "Error converting DOC to DOCX (attempt %s/%s): %s",
                        attempt,
                        max_attempts,
                        stderr.decode("utf-8", errors="ignore"),
                    )
                    if attempt < max_attempts:
                        time.sleep(0.5 * attempt)
                        continue
                    return None

                docx_file = [
                    file
                    for file in os.listdir(temp_dir)
                    if file.endswith(".docx")
                ]
                for file in docx_file:
                    converted_file = os.path.join(temp_dir, file)
                    with open(converted_file, "rb") as f:
                        docx_content = f.read()
                        logger.info(
                            "Successfully read DOCX file, size: %s",
                            len(docx_content),
                        )
                        return docx_content

                logger.warning(
                    "No DOCX produced despite success (attempt %s/%s)",
                    attempt,
                    max_attempts,
                )
                if attempt < max_attempts:
                    time.sleep(0.5 * attempt)
        return None

    def _try_find_executable_path(
        self,
        executable_name: str,
        possible_path: List[str] = [],
        environment_variable: List[str] = [],
    ) -> Optional[str]:
        """查找可执行文件路径。"""
        paths: List[str] = []
        paths.extend(possible_path)
        paths.extend(
            os.environ.get(env_var, "") for env_var in environment_variable
        )
        paths = list(set(paths))

        for path in paths:
            if os.path.exists(path):
                logger.info("Found %s at %s", executable_name, path)
                return path

        result = subprocess.run(
            ["which", executable_name], capture_output=True, text=True
        )
        if result.returncode == 0 and result.stdout.strip():
            path = result.stdout.strip()
            logger.info("Found %s at %s", executable_name, path)
            return path

        logger.warning("Failed to find %s", executable_name)
        return None

    def _try_find_soffice(self) -> Optional[str]:
        """查找 LibreOffice/OpenOffice 可执行文件路径。"""
        possible_paths = [
            "/usr/bin/soffice",
            "/usr/lib/libreoffice/program/soffice",
            "/opt/libreoffice25.2/program/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]
        return self._try_find_executable_path(
            executable_name="soffice",
            possible_path=possible_paths,
            environment_variable=["LIBREOFFICE_PATH"],
        )

    def _try_find_antiword(self) -> Optional[str]:
        """查找 antiword 可执行文件路径。"""
        possible_paths = [
            "/usr/bin/antiword",
            "/usr/local/bin/antiword",
            "C:\\Program Files\\Antiword\\antiword.exe",
            "C:\\Program Files (x86)\\Antiword\\antiword.exe",
        ]
        return self._try_find_executable_path(
            executable_name="antiword",
            possible_path=possible_paths,
            environment_variable=["ANTIWORD_PATH"],
        )


# ============================================================
# ExcelParser（来源：excel_parser.py）
# ============================================================
_IMAGE_FUNC_RE = re.compile(
    r"^=?(_xlfn\.)?(DISPIMG|IMAGE)\(", re.IGNORECASE
)


def _is_image_function(value: object) -> bool:
    """如果 *value* 看起来像嵌入式图片函数字符串，返回 True。"""
    if not isinstance(value, str):
        return False
    return _IMAGE_FUNC_RE.match(value) is not None


class ExcelParser(BaseParser):
    """Excel 文件解析器（.xlsx、.xls）。

    该解析器通过处理所有工作表并将每行转换为结构化文本格式，
    从 Excel 文件中提取文本内容。
    """

    def parse_into_text(self, content: bytes) -> ParsedDocument:
        text: List[str] = []

        excel_file = _open_excel_file(content, file_type=self.file_type)

        for excel_sheet_name in excel_file.sheet_names:
            df = _read_sheet_dataframe(excel_file, excel_sheet_name)
            df.dropna(how="all", inplace=True)

            for _, row in df.iterrows():
                page_content = []
                for k, v in row.items():
                    if pd.notna(v) and not _is_image_function(v):
                        page_content.append("%s: %s" % (k, v))

                if not page_content:
                    continue

                content_row = ",".join(page_content) + "\n"
                text.append(content_row)

        return ParsedDocument(content="".join(text))


def _read_sheet_dataframe(
    excel_file: pd.ExcelFile, sheet_name: str
) -> pd.DataFrame:
    """将工作表读取为具有稳定列标签的 DataFrame。"""
    from openpyxl.utils import (  # type: ignore[import-untyped]  # openpyxl 未安装 stub
        get_column_letter,
    )

    if excel_file.engine == "openpyxl":
        df = excel_file.parse(sheet_name=sheet_name, header=None)
        df.columns = [
            get_column_letter(idx + 1) for idx in range(len(df.columns))
        ]
        return df

    df = excel_file.parse(sheet_name=sheet_name, header=0)
    if df.empty:
        df = excel_file.parse(sheet_name=sheet_name, header=None)
        df.columns = [
            get_column_letter(idx + 1) for idx in range(len(df.columns))
        ]
    elif any(str(col).startswith("Unnamed:") for col in df.columns):
        df = excel_file.parse(sheet_name=sheet_name, header=None)
        df.columns = [
            get_column_letter(idx + 1) for idx in range(len(df.columns))
        ]
    return df


def _prepare_xlsx_bytes(data: bytes) -> bytes:
    repaired = repair_xlsx_bytes(data)
    if repaired is not None:
        data = repaired
    return fill_merged_cells_xlsx(data)


def _open_excel_file(
    content: bytes, file_type: str | None = None
) -> pd.ExcelFile:
    """使用显式引擎选择和回退机制打开 Excel 工作簿。"""
    data = content
    converted_via_soffice = False

    while True:
        ext = detect_excel_format(data)
        if ext is None:
            if converted_via_soffice:
                raise ValueError(
                    "Excel file format cannot be determined, "
                    "you must specify an engine manually."
                )
            try:
                data = normalize_excel_bytes(data, file_type=file_type)
            except ValueError as exc:
                raise ValueError(
                    "Excel file format cannot be determined, "
                    "you must specify an engine manually."
                ) from exc
            converted_via_soffice = True
            continue

        if ext == "ods":
            converted = convert_excel_to_xlsx_bytes(data, suffix=".ods")
            if converted:
                data = converted
                continue

        engine = engine_for_format(ext)
        if ext == "xlsx":
            data = _prepare_xlsx_bytes(data)
            engine = "openpyxl"
        try:
            return pd.ExcelFile(BytesIO(data), engine=engine)
        except ImportError as exc:
            raise ValueError(
                "Excel engine %r is not available for .%s files" % (engine, ext)
            ) from exc
        except KeyError as exc:
            if "sharedStrings.xml" not in str(exc) or engine != "openpyxl":
                raise
            repaired = repair_xlsx_bytes(data)
            if repaired is None:
                raise
            logger.info("Repaired XLSX sharedStrings packaging before parse")
            data = _prepare_xlsx_bytes(repaired)
            continue
        except ValueError as exc:
            if converted_via_soffice or "cannot be determined" not in str(
                exc
            ):
                raise
            try:
                data = normalize_excel_bytes(content, file_type=file_type)
            except ValueError:
                raise
            converted_via_soffice = True
            continue


# ============================================================
# DOCX 解析辅助类型（来源：docx_parser.py）
# ============================================================
class ImageData:
    """表示文档内容中已处理的图片"""

    local_path: str = ""
    object: Optional[Image.Image] = None
    url: str = ""


@dataclass
class LineData:
    """表示文档内容中已处理的一行及其关联图片"""

    text: str = ""
    images: List[ImageData] = field(default_factory=list)
    extra_info: str = ""
    page_num: int = 0
    content_sequence: List[Tuple[str, Any]] = field(default_factory=list)


# ============================================================
# DocxParser — 完整 DOCX 解析器 (src: docx_parser.py)
# ============================================================
class DocxParser(BaseParser):
    """DOCX 文档解析器"""

    def __init__(
        self,
        max_pages: Optional[int] = None,
        **kwargs: Any,
    ):
        """初始化 DOCX 文档解析器

        参数：
            file_name: 文件名
            file_type: 文件类型，如果为 None，则从文件名推断
            enable_multimodal: 是否启用多模态处理
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
            separators: 分隔符列表
            ocr_backend: OCR 引擎类型
            ocr_config: OCR 引擎配置
            max_image_size: 最大图片尺寸限制
            max_concurrent_tasks: 最大并发任务数
            max_pages: 最大处理页数
        """
        super().__init__(**kwargs)
        self.max_pages = (
            settings.docx_max_pages if max_pages is None else max_pages
        )
        if self.max_pages <= 0:
            self.max_pages = 100000  # no limit
        logger.info(
            "DocxParser initialized with max_pages=%s", self.max_pages
        )

    def parse_into_text(self, content: bytes) -> ParsedDocument:
        """解析 DOCX 文档，提取文本内容和图片 Markdown 链接"""
        logger.info(
            "Parsing DOCX document, content size: %s bytes", len(content)
        )
        logger.info("Max pages limit set to: %s", self.max_pages)

        start_time = time.time()
        max_workers = min(4, os.cpu_count() or 2)
        logger.info(
            "Setting max_workers to %s for document processing", max_workers
        )

        try:
            inline_images: Dict[str, str] = {}

            def _inline_upload(local_path: str) -> str:
                """读取临时图片文件，base64 编码，并返回引用路径。"""
                try:
                    with open(local_path, "rb") as f:
                        raw = f.read()
                    ext = os.path.splitext(local_path)[1].lower() or ".png"
                    ref = "images/%s%s" % (uuid7().hex, ext)
                    inline_images[ref] = base64.b64encode(raw).decode()
                    return ref
                except Exception as exc:
                    logger.warning(
                        "Failed to read temp image %s: %s", local_path, exc
                    )
                    return ""

            logger.info(
                "Starting Docx processing with max_pages=%s", self.max_pages
            )
            docx_processor = Docx(
                max_image_size=1920,
                enable_multimodal=True,
                upload_file=_inline_upload,
            )
            all_lines, tables = docx_processor(
                binary=content,
                max_workers=max_workers,
                to_page=self.max_pages,
            )
            processing_time = time.time() - start_time
            logger.info(
                "Docx processing completed in %.2fs, "
                "extracted %s sections and %s tables",
                processing_time,
                len(all_lines),
                len(tables),
            )

            logger.info("Processing document sections")
            section_start_time = time.time()

            text_parts = []
            image_parts: Dict[str, str] = {}

            for sec_idx, line in enumerate(all_lines):
                try:
                    if line.text is not None and line.text != "":
                        text_parts.append(line.text)
                        if sec_idx < 3 or sec_idx % 50 == 0:
                            preview = (
                                line.text[:50] + "..."
                                if len(line.text) > 50
                                else line.text
                            )
                            logger.info(
                                "Added section %s text: %s",
                                sec_idx + 1,
                                preview,
                            )
                    if line.images:
                        for image_data in line.images:
                            if image_data.url and image_data.object:
                                image_parts[image_data.url] = (
                                    endecode.decode_image(image_data.object)
                                )
                                image_data.object.close()
                except Exception as e:
                    logger.error(
                        "Error processing section %s: %s",
                        sec_idx + 1,
                        str(e),
                    )
                    logger.error(
                        "Detailed stack trace: %s", traceback.format_exc()
                    )
                    continue

            section_processing_time = time.time() - section_start_time
            logger.info(
                "Section processing completed in %.2fs",
                section_processing_time,
            )
            logger.info("Combining all text parts")
            text = "\n\n".join([part for part in text_parts if part])

            if not text:
                logger.warning(
                    "Generated text is empty, trying alternative method"
                )
                return self._parse_using_simple_method(content)

            total_processing_time = time.time() - start_time
            logger.info(
                "Parsing complete in %.2fs, generated %s characters of text",
                total_processing_time,
                len(text),
            )

            image_parts.update(inline_images)
            return ParsedDocument(content=text, images=image_parts)
        except Exception as e:
            logger.error(
                "Error parsing DOCX document: %s", str(e)
            )
            logger.error(
                "Detailed stack trace: %s", traceback.format_exc()
            )
            return self._parse_using_simple_method(content)

    def _parse_using_simple_method(
        self, content: bytes
    ) -> ParsedDocument:
        """使用简化方法解析文档，作为回退方案。"""
        logger.info(
            "Attempting to parse document using simplified method"
        )
        start_time = time.time()
        try:
            doc = Document(BytesIO(content))
            logger.info(
                "Successfully loaded document in simplified method, "
                "contains %s paragraphs and %s tables",
                len(doc.paragraphs),
                len(doc.tables),
            )
            text_parts = []

            para_count = len(doc.paragraphs)
            logger.info("Extracting text from %s paragraphs", para_count)
            para_with_text = 0
            for i, para in enumerate(doc.paragraphs):
                if i % 100 == 0:
                    logger.info(
                        "Processing paragraph %s/%s", i + 1, para_count
                    )
                if para.text.strip():
                    text_parts.append(para.text.strip())
                    para_with_text += 1

            logger.info(
                "Extracted text from %s/%s paragraphs",
                para_with_text,
                para_count,
            )

            table_count = len(doc.tables)
            logger.info("Extracting text from %s tables", table_count)
            tables_with_content = 0
            rows_processed = 0
            for i, table in enumerate(doc.tables):
                if i % 10 == 0:
                    logger.info(
                        "Processing table %s/%s", i + 1, table_count
                    )

                table_has_content = False
                for row in table.rows:
                    rows_processed += 1
                    row_text = " | ".join(
                        [
                            cell.text.strip()
                            for cell in row.cells
                            if cell.text.strip()
                        ]
                    )
                    if row_text:
                        text_parts.append(row_text)
                        table_has_content = True

                if table_has_content:
                    tables_with_content += 1

            logger.info(
                "Extracted content from %s/%s tables, processed %s rows",
                tables_with_content,
                table_count,
                rows_processed,
            )

            result_text = "\n\n".join(text_parts)
            processing_time = time.time() - start_time
            logger.info(
                "Simplified parsing complete in %.2fs, "
                "generated %s characters of text",
                processing_time,
                len(result_text),
            )

            if not result_text:
                logger.warning(
                    "No text extracted using simplified method"
                )
                return ParsedDocument()

            return ParsedDocument(content=result_text)
        except Exception as backup_error:
            processing_time = time.time() - start_time
            logger.error(
                "Simplified parsing failed %.2fs: %s",
                processing_time,
                backup_error,
            )
            logger.error(
                "Detailed traceback: %s", traceback.format_exc()
            )
            return ParsedDocument()


class Docx:
    """底层 DOCX 处理器，支持多进程图片提取。"""

    def __init__(
        self,
        max_image_size: int = 1920,
        enable_multimodal: bool = False,
        upload_file: Optional[Callable[[str], str]] = None,
    ) -> None:
        logger.info("Initializing DOCX processor")
        self.max_image_size = max_image_size
        self.picture_cache: dict[str, Any] = {}
        self.enable_multimodal = enable_multimodal
        self.upload_file = upload_file
        self.doc: Any = None  # __call__ 中赋值、结束后置回 None；docx 元素链复杂，Any 显式化

    def get_picture(self, document: Any, paragraph: Any) -> Optional[Image.Image]:
        logger.info("Extracting image from paragraph")
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            logger.info("No image found in paragraph")
            return None
        img = img[0]
        try:
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part = document.part.related_parts[embed]
            logger.info("Found embedded image with ID: %s", embed)

            try:
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logger.warning(
                    "Unrecognized image format. Skipping image."
                )
                return None
            except UnexpectedEndOfFileError:
                logger.warning(
                    "EOF was unexpectedly encountered while reading "
                    "an image stream. Skipping image."
                )
                return None
            except InvalidImageStreamError:
                logger.warning(
                    "The recognized image stream appears to be "
                    "corrupted. Skipping image."
                )
                return None

            try:
                image = Image.open(BytesIO(image_blob)).convert("RGBA")
                logger.info(
                    "Successfully extracted image, size: %sx%s",
                    image.width,
                    image.height,
                )
                return image
            except Exception as e:
                logger.error("Failed to open image: %s", str(e))
                return None
        except Exception as e:
            logger.error("Error extracting image: %s", str(e))
            return None

    def _identify_page_paragraph_mapping(self, max_page: int = 100000) -> dict[int, list[int]]:
        """识别每页包含的段落范围。"""
        start_time = time.time()
        logger.info(
            "Identifying page to paragraph mapping (max_page=%s)", max_page
        )
        page_to_paragraphs: dict[int, list[int]] = {}
        current_page = 0
        page_to_paragraphs[current_page] = []
        total_paragraphs = len(self.doc.paragraphs)
        logger.info("Total paragraphs to map: %s", total_paragraphs)

        if total_paragraphs > 1000:
            logger.info(
                "Large document detected, using heuristic paragraph mapping"
            )
            estimated_paras_per_page = 25
            for p_idx in range(total_paragraphs):
                est_page = p_idx // estimated_paras_per_page
                if est_page > max_page:
                    logger.info(
                        "Reached max page limit (%s) at paragraph %s, "
                        "stopping paragraph mapping",
                        max_page,
                        p_idx,
                    )
                    break
                if est_page not in page_to_paragraphs:
                    page_to_paragraphs[est_page] = []
                page_to_paragraphs[est_page].append(p_idx)
                if p_idx > 0 and p_idx % 1000 == 0:
                    logger.info(
                        "Heuristic mapping: processed %s/%s paragraphs",
                        p_idx,
                        total_paragraphs,
                    )
            mapping_time = time.time() - start_time
            logger.info(
                "Created heuristic mapping with %s pages in %.2fs",
                len(page_to_paragraphs),
                mapping_time,
            )
            return page_to_paragraphs

        logger.info("Using standard paragraph mapping method")
        page_breaks_found = 0
        for p_idx, p in enumerate(self.doc.paragraphs):
            page_to_paragraphs[current_page].append(p_idx)
            if p_idx > 0 and p_idx % 100 == 0:
                logger.info(
                    "Processed %s/%s paragraphs in page mapping",
                    p_idx,
                    total_paragraphs,
                )
            page_break_found = False
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    page_break_found = True
                    break
                if (
                    "w:br" in run._element.xml
                    and 'type="page"' in run._element.xml
                ):
                    page_break_found = True
                    break
            if not page_break_found and p._element.xpath(".//w:sectPr"):
                page_break_found = True
            if page_break_found:
                page_breaks_found += 1
                current_page += 1
                if current_page > max_page:
                    logger.info(
                        "Reached max page limit (%s), stopping page mapping",
                        max_page,
                    )
                    break
                if current_page not in page_to_paragraphs:
                    page_to_paragraphs[current_page] = []
                if page_breaks_found % 10 == 0:
                    logger.info(
                        "Found %s page breaks so far, current page: %s",
                        page_breaks_found,
                        current_page,
                    )

        empty_pages = [
            page
            for page, paras in page_to_paragraphs.items()
            if not paras
        ]
        if empty_pages:
            logger.info(
                "Removing %s empty pages from mapping", len(empty_pages)
            )
            for page in empty_pages:
                del page_to_paragraphs[page]

        mapping_time = time.time() - start_time
        logger.info(
            "Created paragraph mapping with %s pages in %.2fs",
            len(page_to_paragraphs),
            mapping_time,
        )

        if not page_to_paragraphs:
            logger.warning(
                "No valid page mapping created, using fallback method"
            )
            page_to_paragraphs[0] = list(range(total_paragraphs))

        page_sizes = [
            len(paragraphs) for paragraphs in page_to_paragraphs.values()
        ]
        if page_sizes:
            avg_paragraphs = sum(page_sizes) / len(page_sizes)
            min_paragraphs = min(page_sizes)
            max_paragraphs = max(page_sizes)
            logger.info(
                "Page statistics: avg=%.1f, min=%s, max=%s paragraphs per page",
                avg_paragraphs,
                min_paragraphs,
                max_paragraphs,
            )

        return page_to_paragraphs

    def __call__(
        self,
        binary: Optional[bytes] = None,
        from_page: int = 0,
        to_page: int = 100000,
        max_workers: Optional[int] = None,
    ) -> Tuple[MutableSequence[LineData], List[Any]]:
        """处理 DOCX 文档，支持每页的并发处理。"""
        logger.info("Processing DOCX document")

        cpu_count = os.cpu_count() or 2
        logger.info("System has %s CPU cores available", cpu_count)

        self.doc = self._load_document(binary)
        if not self.doc:
            return [], []

        self.para_page_mapping = self._identify_page_paragraph_mapping(
            to_page
        )
        logger.info(
            "Identified page to paragraph mapping for %s pages",
            len(self.para_page_mapping),
        )

        pages_to_process = self._apply_page_limit(
            self.para_page_mapping, from_page, to_page
        )
        if not pages_to_process:
            logger.warning("No pages to process after applying page limits!")
            return [], []

        self._init_shared_resources()
        self._process_document(
            binary,
            pages_to_process,
            from_page,
            to_page,
            max_workers,
        )

        tbls = self._process_tables()
        self.doc = None

        logger.info(
            "Document processing complete, "
            "extracted %s text sections and %s tables",
            len(self.all_lines),
            len(tbls),
        )
        return self.all_lines, tbls

    def _load_document(self, binary: Optional[bytes]) -> Any:
        """从二进制内容加载文档。"""
        try:
            # binary 为 None 时 BytesIO 抛错被 except 捕获并返回 None，cast 仅为类型断言
            # 返回类型 Any：docx.api.Document 是工厂函数不能作类型，实际返回 docx.document.Document
            doc = Document(BytesIO(cast(bytes, binary)))
            logger.info("Successfully loaded document from binary content")
            return doc
        except Exception as e:
            logger.error("Failed to load DOCX document: %s", str(e))
            return None

    def _init_shared_resources(self) -> None:
        """初始化共享资源。"""
        self.lines_lock = threading.Lock()
        # 多进程路径下 all_lines 是 multiprocessing ListProxy，故用 MutableSequence
        self.all_lines: MutableSequence[LineData] = []

    def _apply_page_limit(
        self, para_page_mapping: dict[int, list[int]], from_page: int, to_page: int
    ) -> list[int]:
        """应用页面限制，返回要处理的页面列表。"""
        total_pages = len(para_page_mapping)
        if total_pages > to_page:
            logger.info(
                "Document has %s pages, limiting processing to first %s pages",
                total_pages,
                to_page,
            )
        else:
            logger.info(
                "Document has %s pages, processing all pages (limit: %s)",
                total_pages,
                to_page,
            )

        all_pages = sorted(para_page_mapping.keys())
        pages_to_process = [
            p for p in all_pages if from_page <= p < to_page
        ]

        if pages_to_process:
            logger.info(
                "Will process %s pages from page %s to page %s",
                len(pages_to_process),
                from_page,
                min(
                    to_page,
                    pages_to_process[-1] if pages_to_process else from_page,
                ),
            )
            if len(pages_to_process) < len(all_pages):
                logger.info(
                    "Skipping %s pages due to page limit",
                    len(all_pages) - len(pages_to_process),
                )
            if len(pages_to_process) <= 10:
                logger.info("Pages to process: %s", pages_to_process)
            else:
                logger.info(
                    "First 5 pages to process: %s, last 5: %s",
                    pages_to_process[:5],
                    pages_to_process[-5:],
                )

        return pages_to_process

    def _process_document(
        self,
        binary: Optional[bytes],
        pages_to_process: list[int],
        from_page: int,
        to_page: int,
        max_workers: Optional[int],
    ) -> None:
        """使用多进程处理大型文档。"""
        cpu_count = os.cpu_count() or 2
        doc_contains_images = self._check_document_has_images()

        if max_workers is None:
            max_workers = self._calculate_optimal_workers(
                doc_contains_images, pages_to_process, cpu_count
            )

        temp_file_path = self._prepare_document_sharing(binary)

        args_list = self._prepare_multiprocess_args(
            pages_to_process,
            from_page,
            to_page,
            doc_contains_images,
            temp_file_path,
        )

        self._execute_multiprocess_tasks(args_list, max_workers)
        self._cleanup_temp_file(temp_file_path)

    def _check_document_has_images(self) -> bool:
        """检查文档是否包含图片。"""
        doc_contains_images = False
        if (
            hasattr(self.doc, "inline_shapes")
            and len(self.doc.inline_shapes) > 0
        ):
            doc_contains_images = True
            logger.info(
                "Document contains %s inline images",
                len(self.doc.inline_shapes),
            )
        return doc_contains_images

    def _calculate_optimal_workers(
        self, doc_contains_images: bool, pages_to_process: list[int], cpu_count: int
    ) -> int:
        """计算最佳工作进程数。"""
        if not doc_contains_images or len(pages_to_process) < cpu_count:
            max_workers = min(
                len(pages_to_process), max(1, cpu_count - 1)
            )
        else:
            max_workers = min(len(pages_to_process), cpu_count)
        logger.info(
            "Automatically set worker count to %s", max_workers
        )
        return max_workers

    def _prepare_document_sharing(self, binary: Optional[bytes]) -> str:
        """通过临时文件准备文档共享。"""
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file_path = temp_file.name
        # binary 为 None 时 _load_document 必已失败并提前返回，运行不到此处
        temp_file.write(cast(bytes, binary))
        temp_file.close()
        return temp_file_path

    def _prepare_multiprocess_args(
        self,
        pages_to_process: list[int],
        from_page: int,
        to_page: int,
        doc_contains_images: bool,
        temp_file_path: str,
    ) -> list[tuple[int, list[int], int, int, bool, int, str, bool]]:
        """准备多进程处理的参数列表。"""
        args_list = []
        for page_num in pages_to_process:
            args_list.append(
                (
                    page_num,
                    self.para_page_mapping[page_num],
                    from_page,
                    to_page,
                    doc_contains_images,
                    self.max_image_size,
                    temp_file_path,
                    self.enable_multimodal,
                )
            )
        return args_list

    def _execute_multiprocess_tasks(
        self,
        args_list: list[tuple[int, list[int], int, int, bool, int, str, bool]],
        max_workers: int,
    ) -> None:
        """执行多进程任务。"""
        with Manager() as manager:
            self.all_lines = manager.list()

            logger.info(
                "Processing %s pages using %s processes",
                len(args_list),
                max_workers,
            )

            batch_start_time = time.time()
            with ProcessPoolExecutor(
                max_workers=max_workers
            ) as executor:
                logger.info(
                    "Started ProcessPoolExecutor with %s workers",
                    max_workers,
                )

                future_to_idx = {
                    executor.submit(process_page_multiprocess, *args): i
                    for i, args in enumerate(args_list)
                }
                logger.info(
                    "Submitted %s processing tasks to process pool",
                    len(future_to_idx),
                )

                self._collect_process_results(
                    future_to_idx, args_list, batch_start_time
                )

    def _collect_process_results(
        self,
        future_to_idx: Dict[Future[List[LineData]], int],
        args_list: list[tuple[int, list[int], int, int, bool, int, str, bool]],
        batch_start_time: float,
    ) -> None:
        """收集多进程处理结果。"""
        completed_count = 0
        results = []
        temp_img_paths = set()

        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            page_num = args_list[idx][0]
            try:
                page_lines = future.result()

                for line in page_lines:
                    for image_data in line.images:
                        if (
                            image_data.local_path
                            and image_data.local_path.startswith(
                                "/tmp/docx_img_"
                            )
                        ):
                            temp_img_paths.add(image_data.local_path)

                results.extend(page_lines)
                completed_count += 1

                if completed_count % max(
                    1, len(args_list) // 10
                ) == 0 or completed_count == len(args_list):
                    elapsed_ms = int(
                        (time.time() - batch_start_time) * 1000
                    )
                    progress_pct = int(
                        (completed_count / len(args_list)) * 100
                    )
                    logger.info(
                        "Progress: %s/%s pages processed (%s%%, elapsed: %sms)",
                        completed_count,
                        len(args_list),
                        progress_pct,
                        elapsed_ms,
                    )

            except Exception as e:
                logger.error(
                    "Error processing page %s: %s", page_num, str(e)
                )
                logger.error(
                    "Detailed traceback for page %s: %s",
                    page_num,
                    traceback.format_exc(),
                )

        processing_elapsed_ms = int(
            (time.time() - batch_start_time) * 1000
        )
        logger.info(
            "All processing completed in %sms", processing_elapsed_ms
        )

        self._process_multiprocess_results(results)
        self._cleanup_temp_image_files(temp_img_paths)

    def _process_multiprocess_results(self, results: List[LineData]) -> None:
        """处理多进程结果。"""
        lines = list(results)
        image_upload_start = time.time()

        images_to_process = []
        processed_lines = []
        for i, line_data in enumerate(lines):
            if line_data.images and len(line_data.images) > 0:
                images_to_process.append(i)
                logger.info(
                    "Found line %s with %s images to process",
                    i,
                    len(line_data.images),
                )

        image_url_map = {}
        if images_to_process:
            logger.info(
                "Found %s lines with images to process in main process",
                len(images_to_process),
            )

            for line_idx in images_to_process:
                line_data = lines[line_idx]
                image_paths = line_data.images
                page_num = line_data.page_num

                for image_data in image_paths:
                    if (
                        image_data.local_path
                        and os.path.exists(image_data.local_path)
                        and image_data.local_path not in image_url_map
                    ):
                        try:
                            if not image_data.url:
                                # upload_file 仅在 enable_multimodal 路径被调用，此时必已传入
                                image_url = self.upload_file(  # type: ignore[misc]  # mypy 对 Optional[Callable] 直接调用报 misc
                                    image_data.local_path
                                )
                                if image_url:
                                    image_data.url = image_url
                                    markdown_image = "![](%s)" % image_url
                                    image_url_map[
                                        image_data.local_path
                                    ] = markdown_image
                                    logger.info(
                                        "Added image URL for %s: %s",
                                        image_data.local_path,
                                        image_url,
                                    )
                                else:
                                    logger.warning(
                                        "Failed to upload image: %s",
                                        image_data.local_path,
                                    )
                            else:
                                markdown_image = "![](%s)" % image_data.url
                                image_url_map[
                                    image_data.local_path
                                ] = markdown_image
                                logger.info(
                                    "Using existing URL for image %s: %s",
                                    image_data.local_path,
                                    image_data.url,
                                )
                        except Exception as e:
                            logger.error(
                                "Error processing image from page %s: %s",
                                page_num,
                                str(e),
                            )

            image_upload_elapsed = time.time() - image_upload_start
            logger.info(
                "Finished uploading %s images in %.2fs",
                len(image_url_map),
                image_upload_elapsed,
            )

            for line_data in lines:
                processed_content = []
                if line_data.content_sequence:
                    processed_content = line_data.content_sequence
                    page_num = line_data.page_num

                combined_parts = []
                for content_type, content in processed_content:
                    if content_type == "text":
                        combined_parts.append(content)
                    elif content_type == "image":
                        if (
                            isinstance(content, str)
                            and content in image_url_map
                        ):
                            combined_parts.append(
                                image_url_map[content]
                            )
                        elif (
                            hasattr(content, "local_path")
                            and content.local_path in image_url_map
                        ):
                            combined_parts.append(
                                image_url_map[content.local_path]
                            )

                final_text = "\n\n".join(
                    part for part in combined_parts if part
                )
                processed_lines.append(
                    LineData(
                        text=final_text,
                        page_num=page_num,
                        images=line_data.images,
                    )
                )
        else:
            processed_lines = lines

        sorted_lines = sorted(
            processed_lines, key=lambda x: x.page_num
        )
        self.all_lines = sorted_lines

        logger.info(
            "Finished processing %s lines with interleaved images and text",
            len(self.all_lines),
        )

    def _cleanup_temp_image_files(self, temp_paths: Set[str]) -> None:
        """清理多进程创建的临时图片文件。"""
        if not temp_paths:
            return

        logger.info(
            "Cleaning up %s temporary image files", len(temp_paths)
        )
        deleted_count = 0
        error_count = 0

        for path in temp_paths:
            try:
                if os.path.exists(path):
                    os.unlink(path)
                    deleted_count += 1
                    try:
                        temp_dir = os.path.dirname(path)
                        if temp_dir.startswith(
                            "/tmp/docx_img_"
                        ) and os.path.exists(temp_dir):
                            os.rmdir(temp_dir)
                    except OSError:
                        pass
            except Exception as e:
                logger.error(
                    "Failed to delete temp file %s: %s", path, str(e)
                )
                error_count += 1

        logger.info(
            "Temporary file cleanup: deleted %s, errors %s",
            deleted_count,
            error_count,
        )

    def _cleanup_temp_file(self, temp_file_path: str) -> None:
        """清理临时文件。"""
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(
                    "Removed temporary file: %s", temp_file_path
                )
            except Exception as e:
                logger.error(
                    "Failed to remove temporary file: %s", str(e)
                )

    def _process_tables(self) -> List[tuple[tuple[None, str], str]]:
        """处理文档中的表格。"""
        tbls = []
        table_count = len(self.doc.tables)
        if table_count > 0:
            logger.info("Processing %s tables", table_count)
            for tb_idx, tb in enumerate(self.doc.tables):
                if tb_idx % 10 == 0:
                    logger.info(
                        "Processing table %s/%s",
                        tb_idx + 1,
                        table_count,
                    )
                if len(tb.rows) == 0 or all(
                    len(r.cells) == 0 for r in tb.rows
                ):
                    logger.info("Skipping empty table %s", tb_idx + 1)
                    continue
                table_html = self._convert_table_to_html(tb)
                tbls.append(((None, table_html), ""))
        return tbls

    def _convert_table_to_html(self, table: Any) -> str:
        """将表格转换为 HTML。"""
        html = "<table>"
        for r in table.rows:
            html += "<tr>"
            i = 0
            while i < len(r.cells):
                span = 1
                c = r.cells[i]
                for j in range(i + 1, len(r.cells)):
                    if c.text == r.cells[j].text:
                        span += 1
                        i = j
                i += 1
                html += (
                    "<td>%s</td>" % c.text
                    if span == 1
                    else "<td colspan='%s'>%s</td>" % (span, c.text)
                )
            html += "</tr>"
        html += "</table>"
        return html

    def _safe_concat_images(self, images: List[Image.Image]) -> Optional[Image.Image]:
        """安全地拼接图片列表。"""
        if not images:
            return None
        if len(images) == 1:
            return images[0]

        try:
            logger.info(
                "Attempting to concatenate %s images", len(images)
            )
            total_width = max(
                img.width for img in images if hasattr(img, "width")
            )
            total_height = sum(
                img.height for img in images if hasattr(img, "height")
            )

            if total_width <= 0 or total_height <= 0:
                logger.warning(
                    "Invalid image size, returning the first image"
                )
                return images[0]

            new_image = Image.new(
                "RGBA", (total_width, total_height), (0, 0, 0, 0)
            )
            y_offset = 0
            for img in images:
                if not hasattr(img, "width") or not hasattr(
                    img, "height"
                ):
                    continue
                new_image.paste(img, (0, y_offset))
                y_offset += img.height

            logger.info(
                "Successfully concatenated images, final size: %sx%s",
                total_width,
                total_height,
            )
            return new_image
        except Exception as e:
            logger.error(
                "Failed to concatenate images: %s", str(e)
            )
            logger.error(
                "Detailed error: %s", traceback.format_exc()
            )
            return images[0]


# ============================================================
# Docx 多进程的模块级辅助函数
# ============================================================
def _save_image_to_temp(
    logger: logging.Logger, image: Optional[Image.Image], page_num: int, img_idx: int
) -> Optional[str]:
    """将图片保存到临时文件以在进程间传递。"""
    if not image:
        return None

    try:
        temp_dir = tempfile.mkdtemp(prefix="docx_img_")
        temp_file_path = os.path.join(
            temp_dir, "page_%s_img_%s.png" % (page_num, img_idx)
        )
        image.save(temp_file_path, format="PNG")
        logger.info(
            "[PID:%s] Saved image to temporary file: %s",
            os.getpid(),
            temp_file_path,
        )
        return temp_file_path
    except Exception as e:
        logger.error(
            "[PID:%s] Failed to save image to temp file: %s",
            os.getpid(),
            str(e),
        )
        return None


def process_page_multiprocess(
    page_num: int,
    paragraphs: List[int],
    from_page: int,
    to_page: int,
    doc_contains_images: bool,
    max_image_size: int,
    temp_file_path: Optional[str],
    enable_multimodal: bool,
) -> List[LineData]:
    """专为多进程设计的页面处理函数。"""
    try:
        process_logger = logging.getLogger(__name__)

        if page_num < from_page or page_num >= to_page:
            process_logger.info(
                "[PID:%s] Skipping page %s (out of requested range)",
                os.getpid(),
                page_num,
            )
            return []

        process_logger.info(
            "[PID:%s] Processing page %s with %s paragraphs, "
            "enable_multimodal=%s",
            os.getpid(),
            page_num,
            len(paragraphs),
            enable_multimodal,
        )
        start_time = time.time()

        doc = _load_document_in_process(
            process_logger, page_num, temp_file_path
        )
        if not doc:
            return []

        if not paragraphs:
            process_logger.info(
                "[PID:%s] No paragraphs to process for page %s",
                os.getpid(),
                page_num,
            )
            return []

        combined_text, image_objects, content_sequence = (
            _extract_page_content_in_process(
                process_logger,
                doc,
                page_num,
                paragraphs,
                enable_multimodal,
                max_image_size,
            )
        )

        processed_content = []
        temp_image_index = 0
        image_data_list = []

        if enable_multimodal:
            for i, image_object in enumerate(image_objects):
                img_path = _save_image_to_temp(
                    process_logger, image_object, page_num, i
                )
                if img_path:
                    image_data = ImageData()
                    image_data.local_path = img_path
                    image_data.object = image_object
                    image_data_list.append(image_data)

            process_logger.info(
                "[PID:%s] Saved %s images to temp files for page %s",
                os.getpid(),
                len(image_data_list),
                page_num,
            )

            for content_type, content in content_sequence:
                if content_type == "text":
                    processed_content.append(("text", content))
                else:
                    if temp_image_index < len(image_data_list):
                        processed_content.append(
                            ("image", image_data_list[temp_image_index])
                        )
                        temp_image_index += 1

        line_data = LineData(
            text=combined_text,
            images=image_data_list,
            page_num=page_num,
            content_sequence=processed_content,
        )
        page_lines = [line_data]

        processing_time = time.time() - start_time
        process_logger.info(
            "[PID:%s] Page %s processing completed in %.2fs",
            os.getpid(),
            page_num,
            processing_time,
        )
        return page_lines

    except Exception as e:
        process_logger = logging.getLogger(__name__)
        process_logger.error(
            "[PID:%s] Error processing page %s: %s",
            os.getpid(),
            page_num,
            str(e),
        )
        process_logger.error(
            "[PID:%s] Traceback: %s",
            os.getpid(),
            traceback.format_exc(),
        )
        return []


def _load_document_in_process(
    logger: logging.Logger, page_num: int, temp_file_path: Optional[str]
) -> Any:  # docx.api.Document 是工厂函数不能作类型，实际返回 docx.document.Document
    """在进程中加载文档。"""
    logger.info(
        "[PID:%s] Loading document in process for page %s",
        os.getpid(),
        page_num,
    )
    try:
        if temp_file_path is not None and os.path.exists(temp_file_path):
            doc = Document(temp_file_path)
            logger.info(
                "[PID:%s] Loaded document from temp file: %s",
                os.getpid(),
                temp_file_path,
            )
        else:
            logger.error(
                "[PID:%s] No document source provided", os.getpid()
            )
            return None
        return doc
    except Exception as e:
        logger.error(
            "[PID:%s] Failed to load document: %s", os.getpid(), str(e)
        )
        logger.error(
            "[PID:%s] Error traceback: %s",
            os.getpid(),
            traceback.format_exc(),
        )
        return None


def _extract_page_content_in_process(
    logger: logging.Logger,
    doc: Any,  # docx Document 元素链复杂，Any 显式化
    page_num: int,
    paragraphs: List[int],
    enable_multimodal: bool,
    max_image_size: int,
) -> Tuple[str, List[Any], List[Tuple[str, Any]]]:
    """在进程中提取页面内容。"""
    logger.info(
        "[PID:%s] Page %s: Processing %s paragraphs, "
        "enable_multimodal=%s",
        os.getpid(),
        page_num,
        len(paragraphs),
        enable_multimodal,
    )

    content_sequence: List[Tuple[str, Any]] = []
    current_text = ""
    processed_paragraphs = 0
    paragraphs_with_text = 0
    paragraphs_with_images = 0

    for para_idx in paragraphs:
        if para_idx >= len(doc.paragraphs):
            logger.warning(
                "[PID:%s] Paragraph index %s out of range",
                os.getpid(),
                para_idx,
            )
            continue

        paragraph = doc.paragraphs[para_idx]
        processed_paragraphs += 1

        text = paragraph.text.strip()
        if text:
            cleaned_text = re.sub(r"　", " ", text).strip()
            current_text += cleaned_text + "\n"
            paragraphs_with_text += 1

        if enable_multimodal:
            image_object = _extract_image_in_process(
                logger,
                doc,
                paragraph,
                page_num,
                para_idx,
                max_image_size,
            )
            if image_object:
                if current_text:
                    content_sequence.append(("text", current_text))
                    current_text = ""
                content_sequence.append(("image", image_object))
                paragraphs_with_images += 1

        if processed_paragraphs % 50 == 0:
            logger.info(
                "[PID:%s] Page %s: Processed %s/%s paragraphs",
                os.getpid(),
                page_num,
                processed_paragraphs,
                len(paragraphs),
            )

    if current_text:
        content_sequence.append(("text", current_text))

    logger.info(
        "[PID:%s] Page %s: Completed content extraction, "
        "found %s paragraphs with text, %s with images, "
        "total content items: %s",
        os.getpid(),
        page_num,
        paragraphs_with_text,
        paragraphs_with_images,
        len(content_sequence),
    )

    text_parts = []
    images = []

    for content_type, content in content_sequence:
        if content_type == "text":
            text_parts.append(content)
        else:
            images.append(content)

    combined_text = "\n\n".join(text_parts) if text_parts else ""

    return combined_text, images, content_sequence


def _extract_image_in_process(
    logger: logging.Logger,
    doc: Any,  # docx Document 元素链复杂，Any 显式化
    paragraph: Any,  # docx Paragraph 元素链复杂，Any 显式化
    page_num: int,
    para_idx: int,
    max_image_size: int,
) -> Optional[Image.Image]:
    """在进程中从段落提取图片。"""
    try:
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            return None
        img = img[0]
        logger.info(
            "[PID:%s] Page %s: Found pic element in paragraph %s",
            os.getpid(),
            page_num,
            para_idx,
        )

        try:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                logger.warning(
                    "[PID:%s] Page %s: No embed attribute found in image",
                    os.getpid(),
                    page_num,
                )
                return None
            embed = embed[0]
            if embed not in doc.part.related_parts:
                logger.warning(
                    "[PID:%s] Page %s: Embed ID %s not found in related parts",
                    os.getpid(),
                    page_num,
                    embed,
                )
                return None

            related_part = doc.part.related_parts[embed]
            logger.info(
                "[PID:%s] Found embedded image with ID: %s",
                os.getpid(),
                embed,
            )

            try:
                image_blob = related_part.image.blob
                logger.info(
                    "[PID:%s] Successfully extracted image blob, size: %s bytes",
                    os.getpid(),
                    len(image_blob),
                )
            except Exception as blob_error:
                logger.warning(
                    "[PID:%s] Error extracting image blob: %s",
                    os.getpid(),
                    str(blob_error),
                )
                return None

            try:
                image = Image.open(BytesIO(image_blob)).convert("RGBA")

                if hasattr(image, "width") and hasattr(image, "height"):
                    logger.info(
                        "[PID:%s] Successfully created image object, "
                        "size: %sx%s",
                        os.getpid(),
                        image.width,
                        image.height,
                    )

                    if image.width < 50 or image.height < 50:
                        logger.info(
                            "[PID:%s] Skipping small image (%sx%s)",
                            os.getpid(),
                            image.width,
                            image.height,
                        )
                        return None

                    if (
                        image.width > max_image_size
                        or image.height > max_image_size
                    ):
                        scale = min(
                            max_image_size / image.width,
                            max_image_size / image.height,
                        )
                        new_width = int(image.width * scale)
                        new_height = int(image.height * scale)
                        image = image.resize((new_width, new_height))
                        logger.info(
                            "[PID:%s] Resized image to %sx%s",
                            os.getpid(),
                            new_width,
                            new_height,
                        )

                logger.info(
                    "[PID:%s] Found image in paragraph %s",
                    os.getpid(),
                    para_idx,
                )
                return image
            except Exception as e:
                logger.error(
                    "[PID:%s] Failed to create image from blob: %s",
                    os.getpid(),
                    str(e),
                )
                logger.error(
                    "[PID:%s] Error traceback: %s",
                    os.getpid(),
                    traceback.format_exc(),
                )
                return None
        except Exception as e:
            logger.error(
                "[PID:%s] Error extracting image: %s",
                os.getpid(),
                str(e),
            )
            logger.error(
                "[PID:%s] Error traceback: %s",
                os.getpid(),
                traceback.format_exc(),
            )
            return None
    except Exception as e:
        logger.error(
            "[PID:%s] Error processing image: %s",
            os.getpid(),
            str(e),
        )
        logger.error(
            "[PID:%s] Error traceback: %s",
            os.getpid(),
            traceback.format_exc(),
        )
        return None
